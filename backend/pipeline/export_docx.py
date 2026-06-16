import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from pipeline.transcript_display import get_transcript_segments


def _format_time(value) -> str:
    if isinstance(value, str):
        return value

    try:
        seconds = float(value)
    except (TypeError, ValueError):
        seconds = 0.0

    minutes = int(seconds // 60)
    sec = int(seconds % 60)
    return f"{minutes:02d}:{sec:02d}"


def _meeting_report_sections(result: dict, *, include_empty: bool = False) -> list[dict]:
    report = result.get("meeting_report") or {}
    if not isinstance(report, dict):
        report = {}

    sections = []
    for item in report.get("sections") or []:
        if not isinstance(item, dict):
            continue
        raw_title = str(item.get("title") or "").strip()
        content = str(item.get("content") or "").strip()
        if raw_title or content:
            sections.append({"title": raw_title or "보고서", "content": content})

    content = str(report.get("content") or "").strip()
    if not sections and content:
        sections.append({"title": "회의록 보고서", "content": content})
    if not sections and include_empty:
        sections.append({"title": "회의록 보고서", "content": "보고서 내용이 없습니다."})
    return sections


def _normalized_report_text(value: str) -> str:
    return re.sub(r"\s+", "", value or "")


def _meeting_report_intro(result: dict, sections: list[dict]) -> str:
    report = result.get("meeting_report") or {}
    if not isinstance(report, dict):
        return ""

    content = str(report.get("content") or "").strip()
    if not content or not sections:
        return ""

    section_content = "\n".join(str(section.get("content") or "") for section in sections)
    section_with_titles = "\n".join(
        f"{section.get('title') or '보고서'}\n{section.get('content') or ''}"
        for section in sections
    )
    content_key = _normalized_report_text(content)
    if content_key in {
        _normalized_report_text(section_content),
        _normalized_report_text(section_with_titles),
    }:
        return ""
    return content


def _meeting_report_template_name(result: dict) -> str:
    report = result.get("meeting_report") or {}
    if not isinstance(report, dict):
        report = {}
    snapshot = report.get("templateSnapshot") or report.get("template_snapshot") or {}
    if isinstance(snapshot, dict) and str(snapshot.get("name") or "").strip():
        return str(snapshot.get("name") or "").strip()
    if str(report.get("templateName") or report.get("template_name") or "").strip():
        return str(report.get("templateName") or report.get("template_name")).strip()
    report_template = result.get("report_template") or {}
    report_template_id = str(report.get("templateId") or report.get("template_id") or "").strip()
    if isinstance(report_template, dict) and (not report_template_id or str(report_template.get("id") or "").strip() == report_template_id):
        return str(report_template.get("name") or "").strip()
    return ""


def _report_text_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw_line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        bullet_match = re.match(r"^[-*•]\s+(.+)$", line)
        if bullet_match:
            blocks.append(("bullet", bullet_match.group(1).strip()))
            continue
        numbered_match = re.match(r"^(\d+[.)])\s+(.+)$", line)
        if numbered_match:
            blocks.append(("number", numbered_match.group(2).strip()))
            continue
        blocks.append(("paragraph", line))
    return blocks


def _set_rfonts(r_pr, font_name: str) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def _set_run_font(run, *, size: int = 10, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    _set_rfonts(run._element.get_or_add_rPr(), "맑은 고딕")


def _set_style_font(style, *, size: int, bold: bool = False, color: str | None = None) -> None:
    style.font.name = "맑은 고딕"
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    _set_rfonts(style._element.get_or_add_rPr(), "맑은 고딕")


def _apply_run_font(paragraph, *, size: int = 10) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size=size)


def _add_report_text(doc: Document, text: str) -> None:
    blocks = _report_text_blocks(text)
    if not blocks:
        paragraph = doc.add_paragraph("내용 없음")
        paragraph.paragraph_format.space_after = Pt(6)
        _apply_run_font(paragraph)
        return

    for kind, value in blocks:
        style = "List Bullet" if kind == "bullet" else "List Number" if kind == "number" else None
        paragraph = doc.add_paragraph(value, style=style)
        paragraph.paragraph_format.space_after = Pt(5)
        _apply_run_font(paragraph)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_run_font(run, size=9, bold=bold)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=10)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in (
        ("Heading 1", 15, "1F2937"),
        ("Heading 2", 12, "334155"),
    ):
        style = doc.styles[style_name]
        _set_style_font(style, size=size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(14 if style_name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)


def _add_title_block(doc: Document, title: str, result: dict) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_after = Pt(10)
    title_run = title_paragraph.add_run(title or "회의록")
    _set_run_font(title_run, size=18, bold=True, color="111827")

    rows = [
        ("원본 파일", str(result.get("source_file") or "-")),
        ("처리 일시", str(result.get("created_at") or "-")),
    ]
    if result.get("meeting_purpose"):
        rows.append(("회의 목적", str(result.get("meeting_purpose") or "")))
    report_template_name = _meeting_report_template_name(result)
    if report_template_name:
        rows.append(("보고 양식", report_template_name))

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    for index, (label, value) in enumerate(rows):
        label_cell = table.cell(index, 0)
        value_cell = table.cell(index, 1)
        label_cell.width = Inches(1.15)
        value_cell.width = Inches(5.6)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(label_cell, "F3F6FA")
        _set_cell_text(label_cell, label, bold=True)
        _set_cell_text(value_cell, value)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _add_meeting_report(doc: Document, result: dict, section_no: int, *, include_empty: bool = False) -> int:
    sections = _meeting_report_sections(result, include_empty=include_empty)
    if not sections:
        return section_no

    doc.add_heading(f"{section_no}. 회의록 보고서", level=1)
    section_no += 1
    template_name = _meeting_report_template_name(result)
    if template_name:
        paragraph = doc.add_paragraph(f"보고 양식: {template_name}")
        paragraph.paragraph_format.space_after = Pt(6)
        _apply_run_font(paragraph, size=9)
    intro = _meeting_report_intro(result, sections)
    if intro:
        doc.add_heading("보고서 개요", level=2)
        _add_report_text(doc, intro)
    for index, section in enumerate(sections, start=1):
        doc.add_heading(f"{section_no - 1}.{index} {section.get('title') or '보고서'}", level=2)
        _add_report_text(doc, section.get("content") or "내용 없음")
    return section_no


def export_docx(
    result: dict,
    output_path: str,
    template_path: str = None
) -> str:
    """
    result.json 기반으로 DOCX 파일을 생성한다.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    _configure_document(doc)
        
    summary = result.get("summary", {})
    export_scope = result.get("export_scope") or "full"
    title = summary.get("title", "회의록")
    section_no = 1
    
    _add_title_block(doc, title, result)

    if export_scope == "report":
        _add_meeting_report(doc, result, section_no, include_empty=True)
        doc.save(output_path)
        return output_path

    # 회의 요약
    doc.add_heading(f"{section_no}. 회의 요약", level=1)
    section_no += 1
    doc.add_paragraph(summary.get("overview", "내용 없음"))
    
    # 주요 주제
    doc.add_heading(f"{section_no}. 주요 주제", level=1)
    section_no += 1
    for topic in summary.get("topics", []):
        doc.add_paragraph(topic, style='List Bullet')

    topic_sections = summary.get("topic_sections", []) or []
    if topic_sections:
        doc.add_heading(f"{section_no}. 주제별 내용", level=1)
        section_no += 1
        for section in topic_sections:
            doc.add_heading(section.get("topic", "주제"), level=2)
            if section.get("summary"):
                doc.add_paragraph(section.get("summary"))
            for evidence in section.get("evidence", []) or []:
                doc.add_paragraph(f"근거: {evidence}", style='List Bullet')
            for action in section.get("actions", []) or []:
                doc.add_paragraph(f"할 일: {action}", style='List Bullet')

    participant_summaries = summary.get("participant_summaries", []) or []
    if participant_summaries:
        doc.add_heading(f"{section_no}. 참석자별 요약 AI 초안", level=1)
        section_no += 1
        for participant in participant_summaries:
            doc.add_heading(participant.get("participant", "참석자"), level=2)
            if participant.get("summary"):
                doc.add_paragraph(participant.get("summary"))
            for point in participant.get("key_points", []) or []:
                doc.add_paragraph(f"핵심: {point}", style='List Bullet')
            for action in participant.get("actions", []) or []:
                doc.add_paragraph(f"할 일: {action}", style='List Bullet')
        
    # 결정사항
    doc.add_heading(f"{section_no}. 결정사항", level=1)
    section_no += 1
    for dec in summary.get("decisions", []):
        doc.add_paragraph(dec, style='List Bullet')
        
    # 할 일
    doc.add_heading(f"{section_no}. 할 일", level=1)
    section_no += 1
    for act in summary.get("actions", []):
        doc.add_paragraph(act, style='List Bullet')
        
    # 확인 필요 사항
    doc.add_heading(f"{section_no}. 확인 필요 사항", level=1)
    section_no += 1
    for chk in summary.get("needs_check", []):
        doc.add_paragraph(chk, style='List Bullet')

    if export_scope == "full":
        section_no = _add_meeting_report(doc, result, section_no)

    if export_scope != "organized":
        # 대화록
        doc.add_heading(f"{section_no}. 대화록", level=1)
        segments = get_transcript_segments(result)
        for seg in segments:
            time_str = f"[{_format_time(seg.get('start', 0.0))}]"
            speaker = seg.get("speaker_name") or seg.get("speaker") or ""
            text = seg.get("text", "")

            p = doc.add_paragraph()
            if speaker:
                p.add_run(f"{time_str} {speaker}: ").bold = True
            else:
                p.add_run(f"{time_str} ").bold = True
            p.add_run(text)
        
    doc.save(output_path)
    return output_path
